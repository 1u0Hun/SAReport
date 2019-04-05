from docxtpl import DocxTemplate,RichText,InlineImage
import json
import pymysql
import matplotlib.pyplot as plt
import matplotlib
from docx import Document
from PIL import Image
from docx.shared import Cm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def repoter (json_str):
    data = json.loads(json_str)
    id = data['project_id']
    db = pymysql.connect("localhost","root","toor","sangfor_report1" )
    cursor = db.cursor()
    cursor.execute("select project,producer,production_date,reviewer,start_date,total_num,high_num,mid_num,low_num,risk_level from project where project_id ="+str(id))
    raw_data_1 = cursor.fetchone()
    project = raw_data_1[0]
    producer = raw_data_1[1]
    production_date = raw_data_1[2]
    reviewer = raw_data_1[3]
    start_date = raw_data_1[4]
    total_num = raw_data_1[5]
    high_num = raw_data_1[6]
    mid_num = raw_data_1[7]
    low_num = raw_data_1[8]
    risk_level = raw_data_1[9]

    cursor.execute("select vul_id,sys_id from project_vuln where project_id="+str(id))
    raw_data_2 = cursor.fetchall()
    cursor.execute("select sys_id from project_vuln where project_id="+str(id))
    raw_data_9 = cursor.fetchall()
    sys_list=[]
    for  i in range(0,len(raw_data_9)):
        sys_list.append(raw_data_9[i][0])
    sys_list= list(set(sys_list))

    vul_sort=[]
    for j in range(0,len(sys_list)):
        vul_sort.append([])
        for i in range(0,len(raw_data_2)):
            if raw_data_2[i][1]==sys_list[j]:
                vul_sort[j].append(raw_data_2[i][0])


    vul_and_url=raw_data_2   #漏洞和对应站点
    n = len(raw_data_2)
    vul_id_list=[]   #所有的漏洞
    vul_level_list=[] #漏洞和对应等级
    for i in range(0,n):
        vul_id_list.append(raw_data_2[i][0]) #漏洞列表
        cursor.execute("select vul_level,vul_name,vul_description from vuln where vul_id="+str(raw_data_2[i][0]))
        raw_data_3=cursor.fetchone()
        vul_level_list.append([raw_data_2[i][0],raw_data_3[0],raw_data_3[1],raw_data_3[2]])


    for i in range(0,n):
        cursor.execute("select sys_name from system where sys_id="+str(raw_data_2[i][1]))
        raw_data_8 = cursor.fetchone()
        vul_level_list[i].append(raw_data_8[0])


    vul_show_high=[]  #出现的各个级别漏洞
    vul_show_mid=[]
    vul_show_low=[]
    vul_show=[]           #需要展示的最高级别的漏洞，及其详情
    for i in range(0,n):
        if "高" in vul_level_list[i][1]:
            vul_show_high.append([vul_level_list[i][2],vul_level_list[i][3],vul_level_list[i][4]])
        elif "中" in vul_level_list[i][1]:
            vul_show_mid.append([vul_level_list[i][2],vul_level_list[i][3],vul_level_list[i][4]])
        else:
            vul_show_low.append([vul_level_list[i][2],vul_level_list[i][3],vul_level_list[i][4]])
    if vul_show_high != []:
        vul_show=vul_show_high
        vul_level="高危"
    elif vul_show_mid != []:
        vul_show=vul_show_mid
        vul_level="中危"
    else:
        vul_show=vul_show_low
        vul_level="低危"

    cursor.execute("select participant_id from project_participant where project_id="+str(id))
    raw_data_4= cursor.fetchall()
    participant_list=[]  #
    for i in raw_data_4:
        participant_list.append(i[0])
    participant_info=[]       #参与者信息
    for i in participant_list:
        cursor.execute("select participant_name,telephone from participant where participant_id="+str(i))
        raw_data_5=cursor.fetchone()
        participant_info.append([raw_data_5[0],raw_data_5[1]])

    cursor.execute("select sys_id from project_sys where project_id="+str(id))
    raw_data_6=cursor.fetchall()
    sys_info=[]    #系统信息
    for i in range(0,len(raw_data_6)):
        cursor.execute("select sys_name,sys_url,sys_ip from system where sys_id="+str(raw_data_6[i][0]))
        raw_data_7=cursor.fetchone()
        sys_info.append(raw_data_7)
    docx = DocxTemplate("rule.docx")
    context_project  = {'project':project,'producer':producer,'production_date':production_date,'review':reviewer,
                'total_num':total_num,'high_num':high_num,'mid_sum':mid_num,'low_num':low_num,'start_date':start_date,
                'risk_level':risk_level,'vul_level':vul_level}
    context_participant={'participant_contents': []}
    for i in range(0,len(participant_info)):
        par={'name': participant_info[i][0],'phone': participant_info[i][1]}
        context_participant['participant_contents'].append(par)
    context_project.update(context_participant)

    context_sys={'sys_contents': []}
    for i in range(0,len(sys_info)):
        par={'num':str(i+1),'name': sys_info[i][0],'url': sys_info[i][1],'ip':sys_info[i][2]}
        context_sys['sys_contents'].append(par)
    context_project.update(context_sys)

    context_show={'vul_contents': []}
    for i in range(0,len(vul_show)):
        par={'num':str(i+1),'sys': vul_show[i][2],'name': vul_show[i][0],'abstract':vul_show[i][1],'level':vul_level}
        context_show['vul_contents'].append(par)
    context_project.update(context_show)

    vul_showdata=""
    for i in range(0,len(vul_show)):
        vul_showdata=vul_showdata+str(i+1)+") 网站存在"+vul_show[i][0]+"。"+vul_show[i][1]
        if i+1 != len(vul_show):
            vul_showdata = vul_showdata + "\n"
    rt = RichText(vul_showdata)
    context_vulshow={'vulshow':rt}
    context_project.update(context_vulshow)

    data = {'高危': high_num, '中危': mid_num, '低危': low_num}
    source_data = sorted(data.items(), key=lambda x: x[1], reverse=True)
    labels = [source_data[i][0][:4] for i in range(len(source_data))]  # 设置标签
    fracs = [source_data[i][1] for i in range(len(source_data))]
    explode = [x * 0.01 for x in range(len(source_data))]  # 与labels一一对应，数值越大离中心区越远
    plt.axes(aspect=1)  # 设置X轴 Y轴比例
    matplotlib.rcParams['font.sans-serif']=['SimHei']
    # labeldistance标签离中心距离  pctdistance百分百数据离中心区距离 autopct 百分比的格式 shadow阴 影
    plt.pie(x=fracs,  labels=labels, autopct='%3.1f %%',
            shadow=False, labeldistance=1.1, startangle=0, pctdistance=0.8, center=(-1, 0))
    # 控制位置：bbox_to_anchor数组中，前者控制左右移动，后者控制上下。ncol控制 图例所列的列数。默认值为1。fancybox 圆边
    plt.legend(loc=7, bbox_to_anchor=(1.3, 0.80), ncol=3, fancybox=True, shadow=True, fontsize=8)
    plt.savefig(project+'.jpg')
    img = Image.open(project+'.jpg')
    cropped = img.crop((120, 60, 620, 400))  # (left, upper, right, lower)
    out = cropped.resize((400,270),Image.ANTIALIAS)
    out.save(project+'.jpg')

    context_img={ "image":InlineImage(docx, project+'.jpg')}
    context_project.update(context_img)
    docx.render(context_project)
    docx.save("app/static/download/"+project+"渗透测试报告.docx")

    document = Document("app/static/download/"+project+"渗透测试报告.docx")
    for i in range(0,len(sys_list)):
        document.add_paragraph('4.'+str(i+1)+" "+sys_info[i][0], style='tittle1')
        for j in range(0,len(vul_sort[i])):
            cursor.execute("select vul_name,vul_level,vul_description,vul_suggest,vul_link from vuln where vul_id="+str(vul_sort[i][j]))
            raw_data_10=cursor.fetchone()
            cursor.execute("select sys_url,detail from project_vuln where vul_id="+str(vul_sort[i][j])+ " and sys_id=" +str(sys_list[i]))
            raw_data_11=cursor.fetchone()

            document.add_paragraph('4.'+str(i+1)+"."+str(j+1)+" "+raw_data_10[0], style='tittle2')
            table1=document.add_table(rows=2, cols=1,style="bg")
            table2=document.add_table(rows=2, cols=2,style="bg")
            table3=document.add_table(rows=1, cols=1,style="bg")
            table4=document.add_table(rows=2, cols=2,style="bg")
            table1.rows[0].height=Cm(0.9)
            table1.rows[1].height=Cm(0.9)
            table2.cell(0,0).width=Cm(2.74)
            table2.cell(0,1).width=Cm(13.74)
            table2.cell(1,0).width=Cm(2.74)
            table2.cell(1,1).width=Cm(13.74)
            table2.rows[0].height=Cm(0.9)
            table2.rows[1].height=Cm(0.9)
            table3.rows[0].height=Cm(0.9)
            table4.cell(0,0).width=Cm(2.74)
            table4.cell(0,1).width=Cm(13.74)
            table4.cell(1,0).width=Cm(2.74)
            table4.cell(1,1).width=Cm(13.74)
            table4.rows[0].height=Cm(0.9)
            table4.rows[1].height=Cm(0.9)
            run=table1.cell(0,0).paragraphs[0].add_run("漏洞级别-"+raw_data_10[1])
            run.font.name='宋体'
            run.font.size=Pt(10.5)
            table1.cell(0,0).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run1 = table1.cell(0,1).paragraphs[0].add_run(raw_data_10[0])
            run1.font.name='宋体'
            run1.font.size=Pt(10.5)
            table1.cell(0,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
            run2 = table2.cell(0,0).paragraphs[0].add_run("URL/IP")
            run2.font.name='宋体'
            run2.font.size=Pt(10.5)
            run3 = table2.cell(0,1).paragraphs[0].add_run(raw_data_11[0])
            run3.font.name='宋体'
            run3.font.size=Pt(10.5)
            run4 = table2.cell(1,0).paragraphs[0].add_run("漏洞描述")
            run4.font.name='宋体'
            run4.font.size=Pt(10.5)
            run5 = table2.cell(1,1).paragraphs[0].add_run(raw_data_10[2])
            run5.font.name='宋体'
            run5.font.size=Pt(10.5)
            run6 = table3.cell(0,0).paragraphs[0].add_run("漏洞验证过程：\n")
            run6.font.name='宋体'
            run6.font.size=Pt(10.5)
            #pat = re.compile('>(.*?)<')
            #print(''.join(pat.findall(raw_data_11[1])))

            run7 = table4.cell(0,0).paragraphs[0].add_run("处置建议")
            run7.font.name='宋体'
            run7.font.size=Pt(10.5)
            run8 = table4.cell(0,1).paragraphs[0].add_run(raw_data_10[3])
            run8.font.name='宋体'
            run8.font.size=Pt(10.5)
            run9 = table4.cell(1,0).paragraphs[0].add_run("参考链接")
            run9.font.name='宋体'
            run9.font.size=Pt(10.5)
            run10 = table4.cell(1,1).paragraphs[0].add_run(raw_data_10[4])
            run10.font.name='宋体'
            run10.font.size=Pt(10.5)

    document.add_paragraph("致谢",style="tittle0")
    document.add_paragraph("深信服安全服务团队感谢"+project+"在渗透测试过程中进行沟通、协调的部门和人员,是你们的大力配合，使得我们的工作能够顺利完成。",style="style1")
    document.add_paragraph("了解更多",style="tittle0")
    document.add_paragraph("了解更多安全信息，或关于本文出现的漏洞、攻击方式等详细介绍与建议，可查看深信服安全中心的威胁维基或关注深信服科技公众号了解最新的安全情报。\n    威胁维基：http://sec.sangfor.com.cn/vulns/lst.html\n    公众号及微博：",style="style1")

    pic = document.add_picture('3.jpg')
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.save("app/static/download/"+project+"渗透测试报告.docx")
    filename = project+"渗透测试报告.docx"
    return filename


# data = {"project_id":1}
# data1 = json.dumps(data)
# print(data1)
# repoter(data1)